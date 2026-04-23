"""DOCX rendering for tailored CVs.

Two paths:
1. If templates/cv_template.docx exists, swap Jinja-style {{ placeholders }}
   inside its paragraphs using docxtpl-style substitution (we roll our own
   to avoid adding docxtpl as a dep). Keeps whatever styling the user set.
2. If the template is missing, generate a clean ATS-safe single-column
   document programmatically. No tables, no columns, no graphics.

Output path: reports/{YYYY-MM-DD}/tailored/{company}_{role_slug}.docx
"""

from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from src.models import Job, TailoredCV

log = logging.getLogger("jobhunt.render")

_SLUG_RE = re.compile(r"[^A-Za-z0-9]+")
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_.]*)\s*\}\}")

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "cv_template.docx"
REPORTS_ROOT = Path(__file__).resolve().parent.parent / "reports"


def slug(text: str) -> str:
    s = _SLUG_RE.sub("-", text or "").strip("-").lower()
    return s or "untitled"


def output_path(job: Job, run_date: date) -> Path:
    folder = REPORTS_ROOT / run_date.isoformat() / "tailored"
    folder.mkdir(parents=True, exist_ok=True)
    filename = f"{slug(job.company)}_{slug(job.title)}.docx"
    return folder / filename


def render_cv(
    job: Job,
    profile: dict[str, Any],
    tailored: TailoredCV,
    *,
    run_date: date | None = None,
    template_path: Path | None = None,
    out_path: Path | None = None,
) -> Path:
    """Render a tailored CV to DOCX. Returns the path written."""
    run_date = run_date or date.today()
    out = out_path or output_path(job, run_date)
    tpl = template_path or TEMPLATE_PATH

    if tpl.exists():
        try:
            _render_from_template(tpl, out, profile, tailored)
            log.info("rendered CV via template: %s", out)
            return out
        except Exception as e:  # noqa: BLE001
            log.warning("template render failed (%s); falling back to programmatic", e)

    _render_programmatic(out, profile, tailored)
    log.info("rendered CV programmatically: %s", out)
    return out


# ---- template path ---------------------------------------------------------


def _render_from_template(tpl: Path, out: Path, profile: dict[str, Any], tailored: TailoredCV) -> None:
    doc = Document(str(tpl))
    context = _build_template_context(profile, tailored)
    _replace_placeholders_in_doc(doc, context)
    doc.save(str(out))


def _build_template_context(profile: dict[str, Any], tailored: TailoredCV) -> dict[str, str]:
    ident = profile.get("identity") or {}
    ctx: dict[str, str] = {
        "name": str(ident.get("name", "")),
        "email": str(ident.get("email", "")),
        "phone": str(ident.get("phone", "")),
        "location": str(ident.get("location", "")),
        "linkedin": str(ident.get("linkedin", "")),
        "github": str(ident.get("github", "")),
        "portfolio": str(ident.get("portfolio", "")),
        "summary": tailored.summary,
        "competencies": ", ".join(tailored.competencies_ordered),
    }
    # Per-role blocks (optional; templates may use "experience" as a single block).
    exp_lines: list[str] = []
    for e in tailored.experience:
        exp_lines.append(f"{e.role}, {e.company}")
        for b in e.bullets:
            exp_lines.append(f"  • {b}")
        exp_lines.append("")
    ctx["experience"] = "\n".join(exp_lines).strip()

    edu_lines = []
    for ed in profile.get("education") or []:
        degree = ed.get("degree") or ""
        inst = ed.get("institution") or ""
        start, end = ed.get("start") or "", ed.get("end") or ""
        line = f"{degree}, {inst}".strip(", ").strip()
        if start or end:
            line = f"{line} ({start} to {end})"
        edu_lines.append(line)
    ctx["education"] = "\n".join(edu_lines)

    cert_lines = []
    for c in profile.get("certifications") or []:
        status = c.get("status") or ""
        name = c.get("name") or ""
        issuer = c.get("issuer") or ""
        dt = c.get("date") or ""
        if status == "in-progress":
            cert_lines.append(f"{name} ({issuer}), in progress")
        elif name:
            cert_lines.append(f"{name} ({issuer}{', ' + str(dt) if dt else ''})")
    ctx["certifications"] = "\n".join(cert_lines)

    proj_lines = []
    for pr in profile.get("projects") or []:
        proj_lines.append(f"{pr.get('name', '')}")
        for b in pr.get("bullets") or []:
            proj_lines.append(f"  • {b}")
    ctx["projects"] = "\n".join(proj_lines).strip()

    # Final em-dash scrub. The template author may have used them in layout
    # text — leave those alone since we only touch placeholder substitutions.
    return {k: v.replace("—", ", ") for k, v in ctx.items()}


def _replace_placeholders_in_doc(doc, context: dict[str, str]) -> None:
    def fmt(match: re.Match) -> str:
        key = match.group(1)
        return context.get(key, match.group(0))

    for para in doc.paragraphs:
        _replace_in_paragraph(para, fmt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, fmt)


def _replace_in_paragraph(para, fmt) -> None:
    text = para.text
    if "{{" not in text:
        return
    new_text = _PLACEHOLDER_RE.sub(fmt, text)
    if new_text == text:
        return
    # Collapse runs down to a single run with the substituted text. We lose
    # intra-paragraph formatting but the substituted content preserves the
    # paragraph-level style (heading, bold, etc.), which is what most CV
    # templates actually use.
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


# ---- programmatic fallback -------------------------------------------------


def _render_programmatic(out: Path, profile: dict[str, Any], tailored: TailoredCV) -> None:
    doc = Document()
    _set_default_font(doc)

    ident = profile.get("identity") or {}
    _name_header(doc, ident.get("name", ""))
    _contact_line(doc, ident)

    _section(doc, "Summary")
    doc.add_paragraph(tailored.summary or "")

    if tailored.competencies_ordered:
        _section(doc, "Core Competencies")
        # comma-separated line keeps it ATS-safe (no tables, no columns)
        doc.add_paragraph(", ".join(tailored.competencies_ordered))

    if tailored.experience:
        _section(doc, "Experience")
        profile_roles = {
            ((e.get("company") or "").lower(), (e.get("role") or "").lower()): e
            for e in profile.get("experience") or []
        }
        for role in tailored.experience:
            orig = profile_roles.get((role.company.lower(), role.role.lower()), {})
            header = doc.add_paragraph()
            h1 = header.add_run(f"{role.role}, {role.company}")
            h1.bold = True
            if orig.get("location") or orig.get("start") or orig.get("end"):
                meta_parts = []
                if orig.get("location"):
                    meta_parts.append(str(orig["location"]))
                if orig.get("start") or orig.get("end"):
                    meta_parts.append(f"{orig.get('start','')} to {orig.get('end','')}")
                doc.add_paragraph(" | ".join(meta_parts))
            for bullet in role.bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(bullet)

    edus = profile.get("education") or []
    if edus:
        _section(doc, "Education")
        for ed in edus:
            p = doc.add_paragraph()
            head = f"{ed.get('degree','')}, {ed.get('institution','')}".strip(", ")
            run = p.add_run(head)
            run.bold = True
            meta_bits = []
            if ed.get("start") or ed.get("end"):
                meta_bits.append(f"{ed.get('start','')} to {ed.get('end','')}")
            if ed.get("notes"):
                meta_bits.append(str(ed["notes"]))
            if meta_bits:
                doc.add_paragraph(" | ".join(meta_bits))

    certs = profile.get("certifications") or []
    if certs:
        _section(doc, "Certifications")
        for c in certs:
            name = c.get("name") or ""
            issuer = c.get("issuer") or ""
            status = c.get("status") or "active"
            dt = c.get("date") or ""
            if not name:
                continue
            line = f"{name} ({issuer})" if issuer else name
            if status == "in-progress":
                line = f"{line}, in progress"
            elif dt:
                line = f"{line}, {dt}"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line)

    projs = profile.get("projects") or []
    if projs:
        _section(doc, "Projects")
        for pr in projs:
            if not pr.get("name"):
                continue
            p = doc.add_paragraph()
            p.add_run(pr["name"]).bold = True
            for b in pr.get("bullets") or []:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(str(b))

    doc.save(str(out))


def _set_default_font(doc) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _name_header(doc, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(name or "")
    run.bold = True
    run.font.size = Pt(18)


def _contact_line(doc, ident: dict[str, Any]) -> None:
    bits = [
        ident.get("email"),
        ident.get("phone"),
        ident.get("location"),
        ident.get("linkedin"),
        ident.get("github"),
        ident.get("portfolio"),
    ]
    line = " | ".join(str(b) for b in bits if b)
    if not line:
        return
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run(line)


def _section(doc, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
