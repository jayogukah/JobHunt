"""Tests for the DOCX renderer. Reads back the generated docx to verify
content actually appears in the document, not just that the file exists.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

from src.models import Job, TailoredCV, TailoredExperience
from src.render import output_path, render_cv, slug


PROFILE = {
    "identity": {
        "name": "Praise Ogukah",
        "email": "jogukah@gmail.com",
        "phone": "+234 806 475 6294",
        "location": "Lagos, Nigeria",
        "linkedin": "linkedin.com/in/praise-ogukah",
    },
    "experience": [
        {
            "company": "KPMG West Africa",
            "role": "Senior Automation Developer",
            "location": "Lagos, Nigeria",
            "start": "2022-05",
            "end": "present",
            "bullets": [],
        }
    ],
    "education": [{"institution": "Covenant University", "degree": "B.Eng. Computer Engineering", "start": "2013", "end": "2018"}],
    "certifications": [
        {"name": "UiPath UiARD", "issuer": "UiPath", "date": "2021-12", "status": "active"},
        {"name": "Azure AI-102", "issuer": "Microsoft", "status": "in-progress"},
    ],
    "projects": [{"name": "JobHunt", "bullets": ["Python pipeline for daily discovery."]}],
}

JOB = Job(
    source="greenhouse",
    source_id="1",
    title="Senior AI Automation Engineer",
    company="Example Co",
    location="Remote - EU",
    remote=True,
    description="",
    apply_url="https://example.com/jobs/1",
)

TAILORED = TailoredCV(
    summary="Automation developer with UiPath and Gemini API experience.",
    competencies_ordered=["UiPath", "Gemini API", "Python"],
    experience=[
        TailoredExperience(
            company="KPMG West Africa",
            role="Senior Automation Developer",
            bullets=["UiPath bots for KYC, 1000+ records processed.", "Gemini API for loan underwriting."],
        )
    ],
)


def _read_all(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append(para.text)
    return "\n".join(parts)


def test_slug_sanitises():
    assert slug("Example Co") == "example-co"
    assert slug("Senior AI / Automation Engineer!") == "senior-ai-automation-engineer"
    assert slug("") == "untitled"


def test_output_path_structure(tmp_path: Path, monkeypatch):
    monkeypatch.setattr("src.render.REPORTS_ROOT", tmp_path)
    p = output_path(JOB, date(2025, 4, 23))
    assert p.parent == tmp_path / "2025-04-23" / "tailored"
    assert p.name == "example-co_senior-ai-automation-engineer.docx"


def test_render_programmatic_fallback(tmp_path: Path):
    out = tmp_path / "cv.docx"
    # Point the template path to something that doesn't exist so we hit
    # the programmatic fallback deterministically.
    render_cv(JOB, PROFILE, TAILORED, run_date=date(2025, 4, 23), template_path=tmp_path / "missing.docx", out_path=out)
    assert out.exists()
    body = _read_all(out)
    assert "Praise Ogukah" in body
    assert "jogukah@gmail.com" in body
    assert "Automation developer" in body
    assert "UiPath" in body
    assert "Senior Automation Developer" in body
    assert "KPMG West Africa" in body
    assert "Covenant University" in body
    assert "in progress" in body  # in-progress cert is labelled
    assert "JobHunt" in body


def test_render_from_real_template(tmp_path: Path):
    # Build a minimal docx with placeholders, then render through it.
    tpl = tmp_path / "cv_template.docx"
    src = Document()
    src.add_paragraph("{{ name }}")
    src.add_paragraph("{{ email }} | {{ phone }}")
    src.add_paragraph("Summary: {{ summary }}")
    src.add_paragraph("Stack: {{ competencies }}")
    src.add_paragraph("{{ experience }}")
    src.save(str(tpl))

    out = tmp_path / "cv.docx"
    render_cv(JOB, PROFILE, TAILORED, run_date=date(2025, 4, 23), template_path=tpl, out_path=out)
    body = _read_all(out)
    assert "Praise Ogukah" in body
    assert "jogukah@gmail.com" in body
    assert "UiPath" in body
    # experience block carried through as multi-line text
    assert "Senior Automation Developer" in body
    assert "KPMG West Africa" in body
    # placeholder syntax should be gone
    assert "{{" not in body
